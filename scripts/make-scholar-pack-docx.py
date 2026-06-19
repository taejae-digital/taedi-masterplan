#!/usr/bin/env python3
"""마스터플랜 v0.11.5 학자 학습자료 — DOCX 생성 (deep edition)"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path('/Users/taejae-agent/workspace/taedi-masterplan')
OUT = ROOT / 'archive/pdf/scholar-profiles-learning-materials-v0.11.5-deep.docx'
IMG_DIRS = [ROOT / 'drafts/scholar-images-parallel', ROOT / 'drafts/scholar-images']

TEAL = RGBColor(0x0d, 0x7d, 0x72)
INK = RGBColor(0x11, 0x11, 0x11)
MUTED = RGBColor(0x4b, 0x55, 0x63)

def slug(s):
    return re.sub(r'[^a-z0-9]+', '-', s.lower()).strip('-')

def find_img(en):
    for d in IMG_DIRS:
        if not d.exists():
            continue
        for ext in ('.jpg', '.jpeg', '.png', '.webp'):
            p = d / (slug(en) + ext)
            if p.exists():
                return p
    return None

CITED = [
    {"ko": "에릭 브린욜프슨", "en": "Erik Brynjolfsson", "axis": "O1 경제질서",
     "place": "미국 · Stanford HAI / Stanford Digital Economy Lab",
     "field": "디지털 경제, AI 생산성, 인간-AI 보완",
     "concepts": [
         "AI augmentation: AI는 인간을 대체만 하지 않고 판단·기획·분석을 증폭한다.",
         "task redesign: 직업이 아니라 과업 단위로 재설계해야 생산성이 난다.",
         "intangible assets: 데이터·소프트웨어·조직역량 같은 무형자산이 경쟁력의 중심이 된다."],
     "readings": [
         ("The Second Machine Age", "기계가 육체노동을 넘어서 인지 과업까지 확장하며 경제와 고용 구조를 바꾼다는 큰 그림."),
         ("Machine, Platform, Crowd", "기업 내부의 기계지능, 외부 플랫폼, 대중 협업이 조직의 경계를 다시 정한다."),
         ("Stanford Digital Economy Lab briefs", "AI 생산성 역설, 업무 재설계, 노동시장 전환에 대한 최신 실증 자료.")],
     "know": [
         "마스터플랜 3.1의 출발점은 “AI가 사람을 없앤다”가 아니라 “표준 과업의 비용이 내려가 개인이 더 큰 생산단위를 다룬다”이다.",
         "정체성은 감성어가 아니라, AI가 낮춘 생산비 위에서 무엇을 만들지 정하는 차별화 기준이다.",
         "정책 질문은 일자리 개수보다 AI·데이터·컴퓨팅을 개인이 쓸 수 있는가로 옮겨야 한다."],
     "debate": [
         "생산성 역설: AI 투자만으로 생산성이 오르지 않는 이유는 무엇인가?",
         "자동화와 보완: 어떤 과업은 대체되고 어떤 과업은 인간의 목적 설정을 더 중요하게 만드는가?"],
     "ask": "한국형 공공 AI 인프라는 개인 생산자의 과업 재설계를 어떻게 지원해야 하는가?"},
    {"ko": "다니엘 서스킨드", "en": "Daniel Susskind", "axis": "O1 경제질서",
     "place": "영국 · King’s College London / Oxford Institute for Ethics in AI",
     "field": "기술실업, 일의 미래, 분배와 의미",
     "concepts": [
         "technological unemployment: 기술이 새 일자리보다 기존 노동 수요를 더 빠르게 줄일 수 있다.",
         "work as distribution: 산업사회는 노동을 통해 소득·지위·인정을 배분했다.",
         "meaning after work: 일의 축소는 생계뿐 아니라 사회적 의미의 공백을 만든다."],
     "readings": [
         ("A World Without Work", "자동화가 일자리와 소득 분배, 삶의 의미를 어떻게 흔드는지 다룬 핵심서."),
         ("Growth: A History and a Reckoning", "성장 자체가 만든 번영과 부작용을 함께 보고 다음 경제질서의 기준을 묻는다."),
         ("일의 미래 강연·기고", "기술실업을 단순 복지가 아니라 사회계약 재설계 문제로 설명한다.")],
     "know": [
         "“일자리 보장”은 산업사회 해법이고, AI 시대에는 생산수단 접근·기여 경로·인정 구조가 함께 필요하다.",
         "정체성 기반 생산자 경제는 취미 예찬이 아니라 노동 중심 배분 질서 이후의 대안이다.",
         "새 사회계약은 소득 안전망과 생산 기회 접근을 동시에 다뤄야 한다."],
     "debate": [
         "일이 줄어도 인간은 어떻게 사회적 인정을 얻는가?",
         "기본소득은 충분한가, 아니면 생산수단 접근권까지 가야 하는가?"],
     "ask": "노동 중심 사회계약이 약해질 때 한국은 소득·기여·인정을 어떤 구조로 다시 묶어야 하는가?"},
    {"ko": "마리아나 마추카토", "en": "Mariana Mazzucato", "axis": "O1 경제질서",
     "place": "영국 · UCL Institute for Innovation and Public Purpose",
     "field": "미션경제, 공공가치, 시장 형성 국가",
     "concepts": [
         "market shaping: 국가는 시장 실패를 보정하는 데 그치지 않고 새 시장을 만든다.",
         "public value: 공공투자는 비용이 아니라 사회가 원하는 가치를 조직하는 방식이다.",
         "mission economy: 달 착륙형 미션처럼 공공·민간 역량을 하나의 목표에 맞춘다."],
     "readings": [
         ("The Entrepreneurial State", "인터넷·GPS·바이오 등 핵심 혁신에서 국가의 선도적 역할을 보여준다."),
         ("Mission Economy", "공공목표를 중심으로 조달·규제·투자·민간 협력을 조직하는 프레임."),
         ("IIPP Public Value papers", "공공가치, 공동창조, 미션지향 정책 설계의 실제 도구.")],
     "know": [
         "AI·데이터·컴퓨팅 접근권은 복지 항목이 아니라 미래 생산정책의 기반이다.",
         "국가는 플랫폼 뒤처리자가 아니라 공공 AI 생산수단의 설계자가 되어야 한다.",
         "마스터플랜의 “여건 보장”은 개인에게 결과를 보장하는 것이 아니라 생산 조건을 개방하는 것이다."],
     "debate": [
         "국가가 시장을 만들 때 관료주의와 포획을 어떻게 피하는가?",
         "공공 AI 인프라의 성과는 GDP가 아니라 어떤 공공가치로 측정하는가?"],
     "ask": "태재가 제안할 공공 AI 인프라 미션은 교육·도시·산업 중 어디서 먼저 실험해야 하는가?"},
    {"ko": "대런 아세모글루", "en": "Daron Acemoglu", "axis": "O2 정치질서",
     "place": "미국 · MIT Institute Professor",
     "field": "제도경제학, 정치경제, 기술과 권력",
     "concepts": [
         "inclusive institutions: 번영은 기술보다 권력과 기회가 열린 제도에서 나온다.",
         "directed technological change: 기술의 방향은 자연현상이 아니라 인센티브와 권력이 정한다.",
         "power and progress: 기술 진보가 모두의 진보가 되려면 제도적 통제가 필요하다."],
     "readings": [
         ("Why Nations Fail", "포용적 제도와 착취적 제도가 국가의 장기 성과를 가른다는 대표 논지."),
         ("Power and Progress", "기술이 노동자와 시민에게 이익이 되도록 방향 설정과 권력 조정이 필요하다는 책."),
         ("AI·자동화 논문", "자동화 편향과 인간 보완형 기술의 차이를 실증적으로 논의한다.")],
     "know": [
         "AI 시대 정치는 규제 세부가 아니라 “누가 기술 방향과 이익 배분을 정하는가”의 문제다.",
         "플랫폼·알고리즘은 경제 인프라이면서 시민의 선택지를 정하는 정치 권력이다.",
         "국가는 독점 통치자가 아니라 분산 권력의 오케스트레이터가 되어야 한다."],
     "debate": [
         "AI는 민주주의를 강화할 수도, 기업·국가 권력 집중을 키울 수도 있다. 어느 조건이 차이를 만드는가?",
         "인간 보완형 AI는 시장만으로 충분히 공급되는가?"],
     "ask": "한국은 AI 기술 방향을 민주적으로 조정할 제도 실험을 도시 단위에서 어떻게 만들 수 있는가?"},
    {"ko": "요슈아 벤지오", "en": "Yoshua Bengio", "axis": "O2 정치질서",
     "place": "캐나다 · Université de Montréal / Mila",
     "field": "딥러닝, AI 안전, 프런티어 AI 거버넌스",
     "concepts": [
         "frontier AI risk: 고도 AI는 생산성 기술이면서 권력 집중·안보 위험·통제 실패를 낳을 수 있다.",
         "independent evaluation: 기업 자체평가만으로는 고위험 AI를 검증할 수 없다.",
         "public-interest AI: 고도 AI 연구는 공공성과 안전 기준을 가져야 한다."],
     "readings": [
         ("AI safety 공개서한·인터뷰", "프런티어 AI가 왜 일반 소프트웨어 규제와 다른 위험을 갖는지 설명한다."),
         ("Mila governance 자료", "안전성 연구, 독립 검증, 국제 협력의 제도 필요성을 다룬다."),
         ("딥러닝 강연", "현대 AI 능력 확장의 기술적 배경을 이해하는 기본 자료.")],
     "know": [
         "미·중 AI 경쟁은 산업정책이면서 동시에 안보·안전 질서 문제다.",
         "한국의 포지션은 중립 구호가 아니라 검증 가능한 테스트베드와 사회계약 실험장이다.",
         "접근권 보장과 고위험 통제는 충돌이 아니라 같이 설계할 문제다."],
     "debate": [
         "오픈소스 AI와 안전 통제는 어디까지 양립 가능한가?",
         "국제 AI 안전기구는 권고 수준인가, 감사·제재 권한까지 가져야 하는가?"],
     "ask": "한국이 제안할 AI 안전 검증 도시·산업 테스트베드는 어떤 권한과 데이터를 가져야 하는가?"},
    {"ko": "리처드 플로리다", "en": "Richard Florida", "axis": "O2 도시·정치질서",
     "place": "캐나다 · University of Toronto Rotman School",
     "field": "도시경제, 창조계급, 지역혁신",
     "concepts": [
         "creative class: 창의적 인재와 문화적 다양성이 도시 경쟁력의 핵심이다.",
         "talent-tolerance-technology: 인재·관용·기술이 혁신 도시의 조건이다.",
         "new urban crisis: 창조도시는 혁신과 동시에 주거·불평등 위기를 낳는다."],
     "readings": [
         ("The Rise of the Creative Class", "도시가 창의적 인재를 끌어들이는 조건과 경제 효과를 설명한다."),
         ("The New Urban Crisis", "창조도시 모델의 불평등과 배제 문제를 비판적으로 보완한다."),
         ("도시혁신 강연", "도시 밀도와 다양성이 혁신 네트워크를 만드는 방식.")],
     "know": [
         "AI 시대에도 도시는 사라지지 않는다. 정체성·산업·교육·생활이 가까이 붙는 실험장이 된다.",
         "강소도시는 작은 대도시가 아니라 정체성 훈련과 산업 연결이 가능한 생활권 단위다.",
         "도시 다양성은 문화 구호가 아니라 생산성 조건이다."],
     "debate": [
         "창조계급 전략은 젠트리피케이션을 키우는가?",
         "강소도시는 대도시 집중을 어떻게 보완할 수 있는가?"],
     "ask": "한국의 강소도시는 어떤 산업·문화·교육 조합으로 정체성 기반 생산자를 길러야 하는가?"},
    {"ko": "쇼샤나 주보프", "en": "Shoshana Zuboff", "axis": "O3 사회계약",
     "place": "미국 · Harvard Business School 명예교수",
     "field": "감시자본주의, 플랫폼 권력, 행동 데이터",
     "concepts": [
         "behavioral surplus: 플랫폼은 행동 데이터를 추출해 예측·개입 자산으로 만든다.",
         "instrumentarian power: 명령보다 환경 설계를 통해 행동을 유도하는 권력.",
         "surveillance capitalism: 인간 경험이 데이터 원료가 되는 경제질서."],
     "readings": [
         ("The Age of Surveillance Capitalism", "데이터 추출과 행동 예측 시장이 민주주의와 자유를 어떻게 침식하는지 다룬 대표 저작."),
         ("강연·인터뷰", "감시자본주의의 핵심 개념과 정책적 함의를 요약적으로 파악할 수 있다."),
         ("플랫폼 책임성 논쟁 자료", "추천·노출·평판 권력의 공적 통제 논쟁을 이해한다.")],
     "know": [
         "플랫폼 문제는 개인정보 보호에 그치지 않는다. 선택지·노출·인정 구조를 누가 배열하는가의 문제다.",
         "새 사회계약에는 설명권·거부권·조정권이 들어가야 한다.",
         "정체성 기반 생산자는 플랫폼 노출과 평판 구조 없이는 시장에 도달하기 어렵다."],
     "debate": [
         "감시자본주의 비판은 너무 포괄적인가, 아니면 플랫폼 권력의 본질을 찌르는가?",
         "데이터 소유권만으로 플랫폼 권력을 통제할 수 있는가?"],
     "ask": "마스터플랜의 권리 UI는 추천·노출·평판 권력을 어디까지 설명·거부·조정하게 해야 하는가?"},
    {"ko": "루치아노 플로리디", "en": "Luciano Floridi", "axis": "O3 사회계약",
     "place": "미국/이탈리아 · Yale Digital Ethics Center / University of Bologna",
     "field": "정보윤리, 디지털 거버넌스, infosphere",
     "concepts": [
         "infosphere: 인간은 정보환경 안에서 정체성과 행위를 구성한다.",
         "onlife: 온라인과 오프라인의 구분이 사라진 생활 조건.",
         "digital ethics: 데이터 보호를 넘어 정보환경 설계 윤리로 확장된다."],
     "readings": [
         ("The Fourth Revolution", "인간 이해가 정보기술에 의해 어떻게 바뀌는지 설명한다."),
         ("The Ethics of Information", "정보를 윤리의 기본 단위로 보고 디지털 사회의 규범을 다룬다."),
         ("AI4People / digital governance papers", "AI 윤리 원칙과 거버넌스 설계를 정책 언어로 연결한다.")],
     "know": [
         "디지털 사회계약은 권리 목록이 아니라 정보환경을 어떻게 설계할지의 원칙이다.",
         "정체성은 디지털 정보환경 속에서 형성·발현·인정된다.",
         "AI 윤리는 금지 목록보다 공동체가 어떤 정보환경을 보장할지로 읽어야 한다."],
     "debate": [
         "윤리 원칙은 실제 플랫폼 권력 조정으로 이어질 수 있는가?",
         "infosphere 관점은 국가·도시·가정 경영으로 어떻게 번역되는가?"],
     "ask": "정체성의 형성·발현·인정을 보장하는 정보환경의 최소 조건은 무엇인가?"},
]

CANDIDATES = [
    ("데이비드 오터", "David Autor", "AI·노동", "미국 · MIT",
     "과업 기반 노동시장 분석. AI가 직업 전체보다 과업 묶음을 바꾸며 중간숙련 노동을 재편한다는 실증 근거.", "3.1 경제질서의 노동 재편 실증 보강"),
    ("요하이 벤클러", "Yochai Benkler", "네트워크 생산", "미국 · Harvard",
     "commons-based peer production. 개인·분산 협업·공유지식이 기업 밖 생산을 가능하게 한다.", "정체성 기반 생산자 경제의 협업 구조 보강"),
    ("캐시 오닐", "Cathy O’Neil", "알고리즘 책임", "미국",
     "Weapons of Math Destruction. 점수화·예측모델이 불평등을 자동화하는 위험.", "권리 UI와 알고리즘 설명·거부권 보강"),
    ("케이트 크로퍼드", "Kate Crawford", "AI 권력", "미국/호주 · USC / Microsoft Research",
     "AI는 데이터·노동·광물·전력·군사 인프라 위에 선 물질적 권력 체계라는 비판.", "AI 인프라 접근권과 권력 집중 분석"),
    ("버지니아 유뱅크스", "Virginia Eubanks", "복지 알고리즘", "미국 · University at Albany",
     "자동화된 복지 행정이 빈곤층을 감시·처벌하는 구조를 분석.", "국가 안전망의 알고리즘화 위험 보강"),
    ("메러디스 휘터커", "Meredith Whittaker", "AI 거버넌스", "미국 · Signal Foundation / AI Now",
     "빅테크 집중, 데이터 권력, AI 산업 구조의 민주적 통제 문제.", "플랫폼 권력 조정과 공공 AI 필요성 보강"),
    ("사피야 우모자 노블", "Safiya Umoja Noble", "검색·권력", "미국 · UCLA",
     "검색엔진과 알고리즘이 인종·성별 편견을 재생산한다는 분석.", "노출·추천 구조의 정치성 보강"),
    ("루하 벤저민", "Ruha Benjamin", "기술과 차별", "미국 · Princeton",
     "New Jim Code. 기술이 중립을 가장해 차별을 재생산하는 방식.", "AI 사회계약의 차별 방지 기준 보강"),
    ("제프리 힌턴", "Geoffrey Hinton", "AI 위험", "캐나다/영국 · University of Toronto",
     "딥러닝 개척자. 고도 AI의 통제 실패와 사회적 위험을 경고.", "4.1 AI 안전·안보 논지 보강"),
    ("스튜어트 러셀", "Stuart Russell", "AI 안전", "미국 · UC Berkeley",
     "Human Compatible. 인간 선호와 불확실성을 중심에 둔 안전한 AI 설계.", "프런티어 AI 거버넌스의 기술철학 보강"),
    ("아제이 아그라왈", "Ajay Agrawal", "AI 경제학", "캐나다 · University of Toronto",
     "AI를 예측 비용 하락으로 해석해 의사결정 구조 변화 설명.", "AI 경제질서의 비용 구조 설명 보강"),
    ("카를로타 페레스", "Carlota Perez", "기술혁명", "영국/베네수엘라",
     "기술혁명과 금융자본, 제도 전환의 장기 파동 분석.", "디지털 전환의 장기 역사 프레임 보강"),
    ("사스키아 사센", "Saskia Sassen", "글로벌 도시", "미국 · Columbia",
     "글로벌 도시가 금융·정보·이주 네트워크의 결절점이 되는 방식.", "도시경영과 세계질서 연결 보강"),
    ("에드워드 글레이저", "Edward Glaeser", "도시경제", "미국 · Harvard",
     "도시는 사람을 연결해 학습과 혁신을 증폭하는 장치.", "4.3 강소도시와 학습밀도 보강"),
    ("앨리슨 고프닉", "Alison Gopnik", "아동발달", "미국 · UC Berkeley",
     "아이들은 지시된 수행보다 탐색·놀이 속에서 세계 모델을 만든다.", "4.4 가정경영과 정체성 발견 보강"),
    ("아마르티아 센", "Amartya Sen", "역량접근", "미국/인도 · Harvard",
     "발전은 소득이 아니라 사람이 실제로 할 수 있는 역량과 자유의 확장.", "새 사회계약의 자유·역량 기준 보강"),
]


def set_korean_font(run, size=None, bold=None, color=None):
    run.font.name = 'Apple SD Gothic Neo'
    r = run._element.rPr.rFonts
    r.set(qn('w:eastAsia'), 'Apple SD Gothic Neo')
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def para(doc, text, size=10.5, bold=False, color=None, space_after=4, style=None):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    set_korean_font(run, size=size, bold=bold, color=color)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    set_korean_font(run, size=size)
    p.paragraph_format.space_after = Pt(3)
    return p


def heading(doc, text, level=1, size=None, color=None):
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    set_korean_font(run, size=size, bold=True, color=color or INK)
    return h


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def main():
    doc = Document()
    # 기본 스타일
    style = doc.styles['Normal']
    style.font.name = 'Apple SD Gothic Neo'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Apple SD Gothic Neo')
    style.font.size = Pt(10.5)
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # ===== 표지 =====
    para(doc, 'SCHOLAR LEARNING PACK · DEEP EDITION', size=11, bold=True, color=TEAL, space_after=10)
    para(doc, '마스터플랜 v0.11.5', size=26, bold=True, space_after=2)
    para(doc, '학자 학습자료', size=26, bold=True, space_after=14)
    para(doc, '인용 학자 8명과 후보 학자 16명을 단순 소개가 아니라 대화 준비 자료로 재구성했다. '
              '각 학자는 활동 지역·소속, 핵심 개념, 읽을 자료 요약, 태재팀이 알아야 할 것, 논쟁점, 자문 질문으로 정리했다.',
         size=12, color=MUTED, space_after=20)
    para(doc, '태재미래전략연구원 · 디지털팀 · 2026.06', size=10, color=MUTED, space_after=0)

    doc.add_page_break()

    # ===== 읽는 법 =====
    heading(doc, '읽는 법과 공통 질문', level=1, size=18)
    heading(doc, '읽는 순서', level=2, size=13, color=TEAL)
    bullet(doc, '먼저 8명 인용 학자의 ‘자문 질문’만 훑어 전체 논점을 잡는다.')
    bullet(doc, '그 다음 각 학자의 핵심 개념과 읽을 자료 요약을 연결한다.')
    bullet(doc, '후보 학자 16명은 빠진 논쟁을 보강하는 확장 카드로 본다.')
    heading(doc, '공통 질문', level=2, size=13, color=TEAL)
    bullet(doc, 'AI는 노동을 대체하는가, 생산수단 접근권을 재분배하는가?')
    bullet(doc, '플랫폼·알고리즘 권력은 시장권력인가, 정치권력인가?')
    bullet(doc, '한국은 미·중 AI 경쟁에서 어떤 사회계약 실험장을 제공할 수 있는가?')

    heading(doc, '마스터플랜 축 — 학자 매핑', level=2, size=13, color=TEAL)
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ['마스터플랜 축', '핵심 학자', '주요 논점', '팀 산출물로 번역']
    rows_data = [
        ('3.1 경제질서', '브린욜프슨·서스킨드·마추카토', 'AI 생산성, 노동 이후 분배, 공공 AI 인프라', '정체성 기반 생산자 경제의 조건과 접근권 설계'),
        ('3.2 정치질서', '아세모글루·벤지오·플로리다', '제도와 권력, AI 안전, 도시 다양성', '개인-도시-국가-세계 권력 재배치와 테스트베드'),
        ('3.3 사회계약', '주보프·플로리디', '플랫폼 권력, 정보환경, 권리 UI', '형성·발현·인정 조건과 설명·거부·조정권'),
    ]
    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        set_korean_font(run, size=10, bold=True)
        shade_cell(cell, 'F2F2F2')
    for i, row in enumerate(rows_data, start=1):
        for j, v in enumerate(row):
            cell = table.cell(i, j)
            cell.text = ''
            run = cell.paragraphs[0].add_run(v)
            set_korean_font(run, size=9.5, bold=(j == 0))

    # ===== 인용 학자 8명 =====
    for i, s in enumerate(CITED):
        doc.add_page_break()
        para(doc, s['axis'], size=10, bold=True, color=TEAL, space_after=2)
        heading(doc, f"{i+1}. {s['ko']}  ({s['en']})", level=1, size=18)

        img = find_img(s['en'])
        if img:
            try:
                doc.add_picture(str(img), width=Cm(3.2))
            except Exception:
                pass

        para(doc, f"활동 지역·소속: {s['place']}", size=10.5, bold=True, space_after=2)
        para(doc, f"분야: {s['field']}", size=10.5, color=TEAL, space_after=8)

        heading(doc, '핵심 개념', level=2, size=12, color=TEAL)
        for c in s['concepts']:
            bullet(doc, c)

        heading(doc, '읽을 자료 요약', level=2, size=12, color=TEAL)
        for title, summary in s['readings']:
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(title + ' — ')
            set_korean_font(run, size=10.5, bold=True)
            run2 = p.add_run(summary)
            set_korean_font(run2, size=10.5)
            p.paragraph_format.space_after = Pt(3)

        heading(doc, '태재팀이 알아야 할 것', level=2, size=12, color=TEAL)
        for k in s['know']:
            bullet(doc, k)

        heading(doc, '논쟁점', level=2, size=12, color=TEAL)
        for d in s['debate']:
            bullet(doc, d)

        heading(doc, '자문 질문', level=2, size=12, color=TEAL)
        para(doc, s['ask'], size=12, bold=True, space_after=0)

    # ===== 후보 학자 16명 =====
    doc.add_page_break()
    heading(doc, '후보 학자 16명 — 추가 학습 카드', level=1, size=18)
    para(doc, '인용 학자 8명의 2배수 후보군. 빠진 논쟁을 보강하고, 자문·세미나 초청 후보로 쓴다.',
         size=10.5, color=MUTED, space_after=10)
    for i, (ko, en, area, place, desc, use) in enumerate(CANDIDATES, start=1):
        para(doc, f"{i}. {ko} ({en})", size=12, bold=True, space_after=1)
        para(doc, f"{area} · {place}", size=9.5, color=TEAL, space_after=1)
        para(doc, desc, size=10, space_after=1)
        para(doc, f"마스터플랜 활용: {use}", size=9.5, color=MUTED, space_after=8)

    doc.save(OUT)
    print('DOCX written:', OUT)


if __name__ == '__main__':
    main()
