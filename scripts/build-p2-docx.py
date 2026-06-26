#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
태재 생애주기-공동체 보고서 DOCX 빌더 (흑백, 표 없음)
논지: 공동체는 개인이 권리를 위임하는 계약으로 구현된 실체이며,
      위임은 회수 가능해야 하는데 산업화 시대에 국가로 집중되며 문제가 생겼다.
python-docx로 직접 서식을 구성한다. 색 일절 사용 안 함.
"""
import os
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT = os.path.join(ROOT, "archive", "docx", "p2-lifecycle-community-report.docx")

BODY_FONT = "Noto Serif KR"
BODY_FONT_FALLBACK = "Batang"
HEAD_FONT = "Noto Sans KR"
BLACK = RGBColor(0x00, 0x00, 0x00)

doc = Document()

# ── 페이지 여백 (A4, 학술 여백) ──
for s in doc.sections:
    s.page_width = Mm(210)
    s.page_height = Mm(297)
    s.top_margin = Mm(28)
    s.bottom_margin = Mm(25)
    s.left_margin = Mm(30)
    s.right_margin = Mm(30)

def add_page_numbers(document):
    """모든 섹션 푸터에 가운데 정렬 페이지 번호(현재/전체) 필드 삽입."""
    for section in document.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.text = ""
        def field(instr):
            run = p.add_run()
            b = OxmlElement('w:fldChar'); b.set(qn('w:fldCharType'), 'begin')
            i = OxmlElement('w:instrText'); i.set(qn('xml:space'), 'preserve'); i.text = instr
            e = OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'), 'end')
            run._r.append(b); run._r.append(i); run._r.append(e)
            run.font.size = Pt(9); run.font.name = HEAD_FONT
            run.font.color.rgb = BLACK
            return run
        field("PAGE")
        sep = p.add_run(" / "); sep.font.size = Pt(9); sep.font.name = HEAD_FONT; sep.font.color.rgb = BLACK
        field("NUMPAGES")

def set_run_font(run, name=BODY_FONT, size=10.5, bold=False, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = BLACK
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts'); rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)

def para(text="", size=10.5, bold=False, italic=False, align=None,
         space_before=0, space_after=6, line=1.6, font=BODY_FONT, indent_first=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    if align is not None:
        p.alignment = align
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent_first is not None:
        pf.first_line_indent = Mm(indent_first)
    if text:
        r = p.add_run(text)
        set_run_font(r, name=font, size=size, bold=bold, italic=italic)
    return p

def rich(segments, size=10.5, align=None, space_before=0, space_after=6, line=1.6, indent_first=4.0):  # noqa: E501
    """segments: list of (text, bold) -> 한 문단에 강조 혼합."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.alignment = align if align is not None else WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent_first:
        pf.first_line_indent = Mm(indent_first)
    for text, bold in segments:
        r = p.add_run(text)
        set_run_font(r, name=BODY_FONT, size=size, bold=bold)
    return p

def h1(num, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(f"{num}  {text}" if num else text)
    set_run_font(r, name=HEAD_FONT, size=13.5, bold=True)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '4'); bottom.set(qn('w:color'), '000000')
    pbdr.append(bottom); pPr.append(pbdr)
    return p

def h2(num, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(f"{num}  {text}" if num else text)
    set_run_font(r, name=HEAD_FONT, size=11.5, bold=True)
    return p

def hr():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr'); bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),'single'); bottom.set(qn('w:sz'),'4')
    bottom.set(qn('w:space'),'1'); bottom.set(qn('w:color'),'808080')
    pbdr.append(bottom); pPr.append(pbdr)

# ════════════════════════ 표지 ════════════════════════
for _ in range(3):
    doc.add_paragraph()
para("태재미래전략연구원 · 디지털 시대 마스터플랜 연구",
     size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24, font=HEAD_FONT, indent_first=None)
para("개인의 필요와 공동체\n— 권리 위임의 계약으로 본 시대별 전환 —",
     size=22, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10, line=1.3, font=HEAD_FONT, indent_first=None)
para("생애의 흐름, 노동의 원리, 공동체의 형성·기능, 권리 위임과 회수의 구조",
     size=11.5, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=40, font=BODY_FONT, indent_first=None)
para("태재미래전략연구원", size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4, font=HEAD_FONT, indent_first=None)
para("디지털팀", size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4, font=HEAD_FONT, indent_first=None)
para("2026년 6월", size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, font=BODY_FONT, indent_first=None)
doc.add_page_break()

# ════════════════════════ 초록 ════════════════════════
h1("", "초록")
para(
    "공동체는 개인이 자신의 필요에 의해 권리의 일부를 위임하는 계약을 통해 구현되는 실체다. "
    "개인이 먼저 있고, 생애의 단계마다 달라지는 필요가 공동체를 불러낸다. 위임의 집합이 공동체를 만들고, "
    "더 이상 그 공동체를 필요로 하지 않을 때 개인은 위임한 권리를 거두어들일 수 있어야 한다. "
    "이 연구는 그 '위임과 회수의 계약 논리'를 이론적 토대로 삼아, 농업·산업화·디지털 세 시대의 노동 원리, "
    "생애주기별 필요, 그 필요에 응답하는 공동체의 형성과 역할을 정합적으로 연결한다. "
    "산업화 시대에는 표준 교육과 안정된 일자리라는 그 시대 개인의 필요에 맞추어 권리가 국가로 모였다. "
    "그 집중은 당대의 필요에는 부합했으나, 정체성의 발휘를 원하는 디지털 시대 개인의 필요에는 더 이상 맞지 않는다. "
    "디지털 시대는 그 모인 권리를 가정·이웃·마을·도시·국가·세계 여러 층위로 다시 분산하는 전환으로 해석된다. "
    "이 논의는 홉스·로크·루소의 사회계약론에 뿌리를 두며, 에릭슨의 생애 발달 연구, "
    "브론펜브레너의 생태체계 연구, 오스트롬의 다중심 거버넌스, 센과 누스바움의 역량 이론으로 보완된다. "
    "또한 폴라니의 탈배태·재배태 개념, 마이클 영의 능력주의 비판, 에스핑안데르센의 사회적 투자론, "
    "헥만의 생애 초기 투자 연구가 각 생애단계 논의를 뒷받침한다.",
    space_after=8, indent_first=None)
rich([("키워드  ", True),
      ("권리 위임, 권리 회수, 계약으로서의 공동체, 생애주기, 노동 원리, 디지털 전환, 사회계약론", False)],
     size=10, indent_first=None, space_after=4)

# ════════════════════════ 1. 서론 ════════════════════════
h1("1.", "서론 — 공동체는 어떻게 생겨나는가")
para(
    "이 연구는 공동체를 개인이 자신의 필요에 의해 권리의 일부를 위임하는 계약으로 구현된 실체로 본다. "
    "개인이 먼저 있고, 살아가면서 마주치는 필요가 공동체를 불러낸다. 위임이 철회되면 공동체는 근거를 잃는다. "
    "어떤 부모가 아이를 위해 마을 돌봄 협동조합에 시간과 결정권 일부를 맡기는 것이 곧 위임이며, "
    "이사를 떠나거나 역할을 다하지 못한다고 판단하면 그 위임을 거두어들일 수 있어야 한다는 것이 이 연구의 출발점이다. "
    "이 관점은 공동체를 개인보다 앞서 존재하는 유기체나 자연 질서가 아닌, "
    "개인의 선택과 위임으로 구성·유지·해소되는 계약적 실체로 본다는 점에서 기존의 공동체론과 결을 달리한다.",
    indent_first=None)
para(
    "17세기 영국의 정치철학자 토머스 홉스(Thomas Hobbes, 1588–1679)는 『리바이어던』(1651)에서 "
    "'만인의 만인에 대한 투쟁' 상태를 벗어나기 위해 개인이 권리를 통치자에게 넘긴다는 계약을 상상했다. "
    "그러나 홉스의 모델에서 위임은 사실상 회수 불가능하다. 한번 넘긴 권리는 절대 권력에 귀속된다. "
    "이에 맞서 17세기 영국의 철학자 존 로크(John Locke, 1632–1704)는 『통치론』(1689)에서 "
    "통치자에게 권력을 맡기는 것은 '신탁(trust)'일 뿐이며, 그 믿음이 깨지면 권리를 도로 거두어들일 수 있다고 보았다. "
    "로크의 신탁과 저항권은 이 연구가 말하는 '위임은 회수 가능해야 한다'는 주장의 직접적인 이론적 뿌리다. "
    "18세기 프랑스의 정치철학자 장자크 루소(Jean-Jacques Rousseau, 1712–1778)는 『사회계약론』(1762)에서 "
    "'일반의지(volonté générale)'를 통해 권리가 한 사람이나 한 기관에 집중되어서는 안 된다고 주장했다. "
    "루소는 자치와 분산의 원리를 옹호한 사상가였다. "
    "세 논의를 한 줄로 읽으면 이렇다. 공동체는 개인이 권리를 위임함으로써 만들어지되(홉스), "
    "그 위임은 언제든 회수할 수 있어야 하며(로크), 권리는 어느 한 곳에 집중되어서는 안 된다(루소). 이것이 이 연구의 이론적 뼈대다.")
para(
    "이 뼈대 위에 세 가지 물음을 던진다. 시대마다 달랐던 노동의 원리는 어떻게 개인의 필요와 공동체의 형태를 규정했는가. "
    "생애주기의 단계마다 달라지는 필요에 어떤 공동체가 응답해 왔는가. "
    "디지털 시대의 전환은 왜 권리의 재분산이라는 관점에서 읽혀야 하는가. "
    "이 세 물음에 대한 답이 서로 정합적으로 이어질 때, 우리는 생애주기와 공동체에 관한 하나의 완결된 논리를 갖게 된다.")

# ════════════════════════ 2. 이론적 토대 ════════════════════════
h1("2.", "이론적 토대 — 권리 위임의 계약으로 본 공동체")

h2("2.1", "공동체 = 권리 위임의 계약적 실체")
para(
    "공동체를 권리 위임의 계약으로 보면 일상의 행위들이 모두 위임으로 읽힌다. 개인이 가정에 시간과 돌봄을 바치는 것, "
    "마을 자치 조직에 결정권을 맡기는 것, 국가에 세금을 내고 법을 따르는 것이 모두 권리 위임이다. "
    "공동체의 크기와 권한은 그 공동체에 얼마나 많은 권리가 위임되어 있느냐에 달려 있다. "
    "개인은 혼자서 감당할 수 없는 필요가 생길 때 그 필요를 충족할 공동체에 권리를 위임한다. "
    "필요가 변하면 위임의 대상과 규모도 바뀐다. 청년기에는 경력을 쌓기 위해 도시와 직장에 많은 권리를 위임하지만, "
    "중장년기에는 가정과 이웃으로 위임의 무게중심이 옮겨가고, 노년기에는 돌봄과 존엄을 위한 공동체에 의존이 깊어진다. "
    "이처럼 위임은 생애 전체에 걸쳐 끊임없이 재구성되는 동적 관계다.",
    indent_first=None)
para(
    "노벨 경제학상을 받은 인도 출신 경제학자 아마르티아 센(Amartya Sen, 1933– )은 "
    "발전을 자유의 확장으로 정의했다(『자유로서의 발전』, 1999). 단순한 소득 증가가 아니라 "
    "개인이 실제로 할 수 있는 일과 될 수 있는 존재의 폭이 넓어지는 것이 진정한 발전이다. "
    "미국의 철학자 마사 누스바움(Martha Nussbaum, 1947– )은 이를 '역량(capabilities)'—실제로 실행할 수 있는 능력—으로 "
    "구체화했다(『역량의 창조』, 2011). 공동체가 개인에게 역량을 제공하는 한에서만 위임은 정당성을 갖는다.")

h2("2.2", "위임은 회수 가능해야 한다 — 산업화 시대 권리 집중의 문제")
para(
    "로크의 신탁 이론이 가리키는 것은 위임이 조건부라는 점이다. 국가가 그 권리를 올바르게 사용하리라는 믿음이 깨지면 "
    "권리는 다시 개인에게로 돌아와야 한다. 그러나 산업화 시대에는 이 회수의 회로가 막혔다. "
    "국가는 교육·고용·복지·의료를 중앙에서 관리하기 시작했고, 개인은 태어나면 국가 교육에 들어가고 "
    "졸업하면 국가가 인증한 기업에 취직하며 노년에는 국가 연금으로 생존하는 단선적 경로에 묶였다. "
    "이 경로 바깥에 서는 것은 곧 사회적 낙오를 의미했다. "
    "장애인이나 소수 언어 화자처럼 표준 경로가 맞지 않는 사람은 제도 안에서 보이지 않는 사람이 되고 만다.",
    indent_first=None)
para(
    "이 과정에서 이웃과 마을은 공동 돌봄의 기능을 잃고 해체되었으며, 가정은 점점 좁은 역할만 남게 되었다. "
    "반면 국가는 거의 모든 기능을 빨아들이며 비대해졌다. 권리가 집중된 공동체는 다양한 개인의 필요에 맞춤 대응할 수 없다. "
    "표준화는 평균 바깥의 개인을 배제하고, 획일화는 다양성을 억압한다.")
para(
    "헝가리 출신 경제사학자 칼 폴라니(Karl Polanyi, 1886–1964)는 『거대한 전환』(1944)에서 "
    "산업화 시대의 시장 경제가 노동·토지·화폐를 사회적 관계에서 뜯어내어 상품으로 만들었다고 분석했다. "
    "폴라니는 이를 '탈배태(disembedding)'—사람들이 원래 속해 있던 관계의 그물에서 떨어져 나오는 것—으로 불렀다. "
    "산업화 시대는 이웃·마을·가정에 분산되어 있던 권리 위임을 국가 하나로 집중시킨 시대였다. "
    "디지털 시대의 과제는 그 집중을 여러 층위로 되돌리는 것, 곧 폴라니가 말한 '재배태(re-embedding)'이다.")

h2("2.3", "다층 자치와 생태체계 — 권리 분산의 근거")
para(
    "여성 최초로 노벨 경제학상을 받은 미국의 정치경제학자 엘리너 오스트롬(Elinor Ostrom, 1933–2012)은 "
    "전 세계 사례를 조사하며 사람들이 중앙 정부나 사유화된 시장 없이도 스스로 자치 규칙을 만들어 "
    "공유 자원을 지속 가능하게 관리할 수 있다는 것을 실증했다(『공유의 비극을 넘어』, 1990). "
    "그녀는 이를 '다중심 거버넌스(polycentric governance)'—여러 층위의 공동체가 각자의 영역에서 자치적으로 문제를 나누어 맡는 구조—라고 불렀다. "
    "권리는 가정·이웃·마을·도시·국가·세계가 층위에 따라 나누어 맡아야 하며, "
    "어느 하나가 모든 권리를 독점하면 나머지 층위는 고사하고 개인은 선택지를 잃는다.",
    indent_first=None)
para(
    "미국의 발달심리학자 유리 브론펜브레너(Urie Bronfenbrenner, 1917–2005)는 개인의 발달을 "
    "여러 겹으로 중첩된 생태체계 안에서 이해했다(『인간 발달의 생태학』, 1979). "
    "가정(미시체계), 이웃과 학교(중간체계), 부모의 직장(외체계), 사회의 가치 규범(거시체계)이 중첩되어 상호작용하는 구조다. "
    "그의 핵심 통찰은 개인에게 가장 깊은 영향을 미치는 것이 가까운 공동체와의 직접적 상호작용—'근위 과정(proximal process)'—이라는 점이다. "
    "산업화 시대는 이 근위 과정이 일어나야 할 가정·이웃·마을을 약화시키고 거시체계인 국가에 권리를 집중함으로써 "
    "개인의 발달을 가장 중요한 토대에서 분리시켰다.")

# ════════════════════════ 3. 노동의 원리와 공동체 ════════════════════════
h1("3.", "세 시대의 노동 원리 — 협동, 분업, 정체성 발휘")

h2("3.1", "농업 시대의 노동 원리: 협동")
para(
    "농업 시대의 노동 원리는 협동이었다. 혼자서는 감당할 수 없는 땅 갈기·씨 뿌리기·수확·방어를 함께 해야 "
    "생존할 수 있었다. 가정에는 양육과 출산의 역할이, 이웃에는 품앗이와 두레의 노동이, "
    "마을에는 공동 경작지 관리와 자치 의사결정이 위임되었다. "
    "국가는 신분 질서와 조세를 관장하는 먼 존재로 개인의 일상에 깊이 개입하지 않았다. "
    "한국의 품앗이—이웃집 모내기를 함께 해주고 추수 때 다시 도움을 받는 방식—는 "
    "비공식적이지만 정교한 권리 위임과 회수의 순환을 보여주는 사례다. "
    "생애주기도 협동의 논리를 따랐다. 아이는 공동체 노동에 참여하기 위한 준비를 했고, "
    "청년기는 혼인과 토지 상속으로 새 가정을 구성하는 시기였으며, "
    "중장년은 마을 자치의 중심이었고, 노년은 권위와 봉양의 대상으로 공동체에 머물렀다.",
    indent_first=None)

h2("3.2", "산업화 시대의 노동 원리: 분업")
para(
    "산업화 시대의 노동 원리는 분업이었다. 복잡한 생산 과정을 잘게 쪼개어 각자에게 반복 담당하게 하는 방식으로 "
    "공장과 관료제 조직은 생산성을 획기적으로 높였다. 한 사람이 전체를 이해할 필요 없이 "
    "정해진 직무를 정확히 수행하는 표준화된 노동력이 필요했다. "
    "'좋은 일자리'를 위한 표준화된 교육, 입시 경쟁, "
    "정규직 고용과 연금으로 이어지는 단선 경로가 삶의 목표가 되었다.",
    indent_first=None)
para(
    "분업의 논리는 공동체에도 적용되었다. 국가는 교육·고용·복지·의료를 일괄 관리하는 거대한 분업 기구가 되었고, "
    "이웃과 마을은 기능을 잃었다. 개인은 여러 공동체에 권리를 분산해 위임하던 상태에서 "
    "국가와 도시 두 층위에만 의존하는 상태로 바뀌었다. "
    "영국의 사회학자 마이클 영(Michael Young, 1915–2002)은 『능력주의의 부상』(1958)에서 "
    "이런 시스템이 탈락자에게 '네가 못난 것'이라는 낙인을 찍으며 "
    "구조적 배제를 개인의 실패로 위장한다고 경고했다.")

h2("3.3", "디지털 시대의 노동 원리: 정체성 발휘")
para(
    "디지털 시대의 노동 원리는 정체성 발휘다. 인공지능과 자동화가 반복 직무를 대체하면서 "
    "표준화된 분업 노동의 경제적 가치는 낮아지고, 각 개인만이 가진 고유한 역량의 가치는 높아진다. "
    "예를 들어 세금 계산서 처리는 자동화로 대체되지만, "
    "복잡한 세무 분쟁을 풀어내는 맥락 이해력은 오히려 더 귀해진다. "
    "플랫폼과 디지털 도구가 기획·제작·유통의 비용을 낮추면서 개인은 큰 조직에 들어가야만 "
    "생산자가 되는 제약에서 풀려나고 있다. 일한다는 것은 조직이 부여한 직무를 수행하는 것이 아니라 "
    "자신의 고유한 정체성을 실제 기여로 발휘하는 것이다.",
    indent_first=None)
para(
    "이 전환은 권리 위임의 방식도 바꾼다. 국가 하나에 집중되었던 권리가 다시 여러 공동체로 분산될 조건이 만들어진다. "
    "가정은 정체성의 단서를 가장 먼저 읽는 공간으로 중요성이 되살아나고, "
    "이웃과 마을은 정체성을 실험하고 훈련하는 관계망으로 복원된다. "
    "도시는 그 정체성을 실제 기여로 발휘하는 무대가 되며, "
    "국가는 모든 것을 관리하는 기관에서 여러 공동체가 잘 작동하도록 최소 조건을 보장하는 조율자로 역할이 바뀐다. "
    "세계는 처음으로 개인의 발휘 무대로 직접 열린다. 국경을 넘는 협업이 가능해지면서, "
    "국내 노동시장에서 자리를 얻지 못한 사람도 세계 무대에서 기여할 경로가 생긴다.")

# ════════════════════════ 4. 생애주기별 필요와 공동체의 응답 ════════════════════════
h1("4.", "생애주기별 필요 — 어떤 공동체에 무엇을 위임하는가")

h2("4.1", "태아·영아 — 단서가 드러나는 시기")
para(
    "태아·영아기는 고유한 가능성의 단서가 처음 드러나는 시기다. 농업 시대의 필요는 무사한 출산과 건강한 몸이었고, "
    "가정·이웃·마을이 공동으로 출산을 돕고 아이를 거두었다. 산업화 시대에는 표준 발달표에 맞는 보살핌이 필요로 규정되었고, "
    "국가가 임신·출산·영유아 건강 관리를 담당하며 가정을 보완했다. "
    "디지털 시대에는 타고난 기질과 강점의 단서를 일찍 발견하고 기록하는 것이 새로운 필요로 등장한다. "
    "생체·행동 데이터를 읽는 기술이 등장하면서, 아이가 어떤 방식으로 환경에 반응하는지를 예전보다 훨씬 일찍 포착할 수 있게 되었다. "
    "가정의 역할은 단순 양육에서 정체성의 단서를 가장 먼저 읽고 기록하는 생활 기반으로 깊어진다. "
    "미국의 경제학자 제임스 헥만(James Heckman, 1944– )은 장기 추적 연구에서 "
    "생애 초기 양육·교육 환경에 대한 투자가 이후 어느 시기보다 높은 복리 효과를 낳는다는 것을 실증했다. "
    "돌봄의 질을 높이는 데 1달러를 쓰면 성인이 된 후 범죄 예방·소득 향상·복지 지출 감소로 수배의 사회적 수익이 돌아온다. "
    "국가는 표준 관리자에서 '먼저 찾아가는 보장자'로 역할을 전환해야 한다.",
    indent_first=None)

h2("4.2", "유아·아동 — 발견하고 실험하는 시기")
para(
    "유아·아동기는 여러 강점의 단서를 체험으로 확인하는 시기다. 농업 시대에는 집안일과 농사일을 익히는 것이, "
    "산업화 시대에는 취학 준비와 기초 학력 습득이 과제였다. 두 시대 모두 아이의 필요는 외부 사회가 정의했다. "
    "디지털 시대에는 다양한 체험으로 관심과 강점을 발견하고 실험하는 것이 이 단계의 과제가 된다. "
    "한 아이는 블록 쌓기에서 공간 감각의 단서를 보이고, 다른 아이는 친구의 감정을 빠르게 읽는 공감 능력을 드러낸다. "
    "이를 위해 가정은 관심을 지지하고, 이웃은 다양한 체험 기회를 열며, 마을은 공동 발견 프로그램을 제공한다. "
    "가정 하나 혹은 국가 하나가 교육을 독점하는 구조에서 벗어나 가정·이웃·마을이 나누어 역할을 맡는 "
    "분산형 위임 구조로 이행하는 것이 이 단계의 핵심이다. "
    "브론펜브레너가 중간체계(mesosystem)라고 부른 가정·이웃·마을의 상호 강화 구조가 이 단계에서 결정적으로 작동해야 한다. "
    "체험이 풍부한 아이와 그렇지 못한 아이의 간격은 이미 이 단계에서 벌어지기 시작하므로, "
    "국가는 기회의 바닥을 보장하는 역할로 이 분산 구조를 뒤에서 떠받쳐야 한다.",
    indent_first=None)

h2("4.3", "청소년 — 훈련하고 방향을 정하는 시기")
para(
    "청소년기는 정체성의 위기가 가장 첨예하게 나타나는 시기다. "
    "독일 태생의 미국 발달심리학자 에릭 에릭슨(Erik H. Erikson, 1902–1994)은 생애 발달을 여덟 단계로 나누어 "
    "각 단계의 심리사회적 과제를 제시했다(『아동기와 사회』, 1950). "
    "청소년기의 핵심 과제는 '정체성 대 역할 혼미'였다. 나는 누구인가, 무엇을 할 수 있는가를 탐색하고 방향을 정하는 것이 이 단계의 본질적 필요다.",
    indent_first=None)
para(
    "농업 시대에는 가업 기술 습득이, 산업화 시대에는 입시 경쟁과 상급 학교 진학이 이 필요를 채웠다. "
    "성적이 정체성을 규정했고, 도시의 입시 학교와 국가의 대입 제도가 청소년의 권리 위임을 거의 독점했다. "
    "마이클 영이 경고했듯 이 구조는 탈락자에게 구조적 배제를 개인의 실패로 위장했다. "
    "디지털 시대에는 마을이 멘토와 훈련 경로를 제공하고, 도시가 심화 경로와 재진입 통로를 마련하여 "
    "각자의 강점이 점수 서열 바깥에서 길러질 수 있는 공간이 만들어진다.")

h2("4.4", "청년 — 발휘하는 시기")
para(
    "청년기는 훈련을 통해 다듬어진 강점을 실제 기여로 발휘하는 시기다. "
    "농업 시대에는 토지와 혼인을 통해 새 가정을 구성하는 것이, 산업화 시대에는 좋은 직장에 취업하고 정착하는 것이 목표였다. "
    "두 경우 모두 필요는 외부가 규정했고 국가와 도시가 그 경로를 통제했다.",
    indent_first=None)
para(
    "디지털 시대의 청년이 필요로 하는 것은 자기 정체성을 실제 기여로 발휘할 무대다. "
    "인공지능과 플랫폼이 진입 비용을 낮추면서 개인은 대기업에 고용되지 않고도 가치를 생산할 수 있게 되었다. "
    "산업화 시대에 청년은 노동력과 결정권을 기업과 국가에 대규모로 위임하는 것이 거의 유일한 선택지였다면, "
    "디지털 시대에는 도시에 발휘의 기반을 위임하고 세계와 선택적으로 협업하며 언제든 재편할 수 있는 "
    "유연한 위임 구조가 가능해진다. "
    "덴마크 출신의 사회정책학자 예스타 에스핑안데르센(Gøsta Esping-Andersen, 1947– )은 "
    "지식경제 시대에는 복지를 소득 보전에만 쓰지 말고 교육·보육·재훈련 같은 사람에 대한 사전 투자—"
    "'사회적 투자(social investment)'—로 전환해야 한다고 강조했다.")

h2("4.5", "중장년 — 재설계하고 확장하는 시기")
para(
    "중장년기는 디지털 전환의 혜택과 위험이 가장 불균등하게 갈리는 시기다. "
    "산업화 시대의 목표는 고용 유지와 가족 부양이었고, 권리 위임은 주로 직장과 국가를 향해 있었다. "
    "안정된 고용이 유지되는 한 개인은 자신의 노동 시간과 역량을 기업에 위임하고 연금 권리는 국가에 맡겨두는 계약을 따랐다. "
    "디지털 전환은 자동화와 인공지능이 기존 직무를 대체하면서 재학습과 강점의 재설계를 요구한다. "
    "마을과 도시가 재학습의 접근로를 열고 가정과 국가가 전환기 보호를 제공하는 것은, "
    "빼앗겼던 선택권을 다시 분산해 돌려주는 과정이다. "
    "직업 전환을 원하는 중장년이 마을 단위 재교육 센터에서 새 기술을 배우고 "
    "국가는 그 기간 동안 소득을 일부 보전해주는 구조가 이에 해당한다.",
    indent_first=None)
para(
    "에릭슨은 중장년기의 핵심 과제를 '생산성 대 침체(generativity vs. stagnation)'로 정의했다. "
    "자신이 쌓은 것을 새 방향으로 확장하고 다음 세대를 위한 무언가를 만들어내지 못하면 정체와 고립으로 빠진다. "
    "이 심리적 위기에 응답하는 것은 개인의 의지가 아니라 공동체의 구조다. "
    "재학습의 기회를 만들고 새 역할을 연결하며 전환기의 불안을 함께 감당하는 공동체가 있어야 한다.")

h2("4.6", "노년 — 전수하고 연결하는 시기")
para(
    "농업 시대의 노년은 권위와 봉양, 산업화 시대의 노년은 은퇴와 부양의 시기였다. "
    "두 시대 모두 노인을 '받는 존재'로 정의했다. "
    "디지털 시대는 건강을 미리 예측·예방하면서 존엄과 판단의 주도권을 지키고, "
    "경험을 다음 세대에 전수하며 세대를 연결하는 '능동적 전수자'로서의 노년을 제안한다.",
    indent_first=None)
para(
    "에릭슨은 노년기의 과제를 '자아통합 대 절망(ego integrity vs. despair)'으로 보았다. "
    "자신의 삶 전체를 돌아보며 '그래도 이 삶은 의미가 있었다'는 통합감을 얻지 못하면 절망에 빠진다. "
    "전수할 통로가 있고 역할이 주어지며 관계 안에 머물 수 있는 공동체 구조가 통합을 가능하게 한다. "
    "가정은 돌봄과 존엄의 생활 공간을, 이웃은 관계의 지속을, 마을과 도시는 멘토와 자문의 역할을 연결하고, "
    "국가는 소득의 최소 기반과 고립 방지의 역할을 맡는다. "
    "폴라니의 언어로 말하면, 노인이 시장과 국가 연금에만 의존하는 상태는 탈배태의 전형이다. "
    "재배태은 노인이 관계와 역할의 그물 안으로 돌아오는 것이며, 그 그물을 유지하는 것은 "
    "각 공동체가 노인으로부터 위임받은 관계·시간·경험을 제대로 활용하고 돌려주는 것이다.")

# ════════════════════════ 5. 공동체의 형성과 역할 ════════════════════════
h1("5.", "공동체의 형성·역할·기능 — 층위별 권리 위임의 구조")

h2("5.1", "가정 — 정체성의 첫 번째 공동체")
para(
    "가정은 개인이 태어나는 순간부터 가장 먼저 권리를 위임받는 공동체다. "
    "농업 시대에 가정은 노동·양육·부양·교육을 모두 떠안는 전능한 공동체였고, "
    "산업화 시대에는 교육을 학교에, 의료를 병원에 위임하면서 역할이 좁아졌다. "
    "디지털 시대에 가정의 핵심 역할은 정체성의 단서를 가장 먼저 읽는 생활 기반이다. "
    "아이가 어디에 반응하는지, 무엇을 할 때 활기를 띠는지를 가장 가까이에서 관찰하고 기록하는 것이다. "
    "또한 중장년의 전환기와 노년의 돌봄·존엄을 지지하는 역할도 함께 맡는다. "
    "브론펜브레너가 말한 미시체계(microsystem)—개인이 직접 몸담는 가장 가까운 환경—가 바로 가정이며, "
    "이 층위가 건강해야 다른 모든 공동체에서의 위임도 의미를 가질 수 있다.",
    indent_first=None)

h2("5.2", "이웃 — 실험의 공동체")
para(
    "이웃은 산업화 시대에 가장 크게 위축된 공동체다. 표준 제도가 모든 것을 조직하면서 "
    "이웃 간 상호 돌봄과 협동의 기능은 사라졌다. 정체성을 발휘하는 시대에 이웃은 다시 중요해진다. "
    "아이가 흥미를 실험하고 어른이 새 관계를 맺는 일상망이다. "
    "관심 기반 모임, 생활 멘토, 상호 돌봄의 품앗이가 이웃에서 이루어질 때 개인은 발견의 기회를 얻는다. "
    "단절된 사람이 다시 관계망으로 돌아오는 첫 번째 통로도 이웃이다. "
    "오스트롬이 연구한 스위스 알프스의 목초지 공동 관리, 일본 어촌의 어장 자치는 "
    "이웃 규모의 자치가 얼마나 정교하게 작동할 수 있는지를 보여준다.",
    indent_first=None)

h2("5.3", "마을 — 훈련과 재학습의 공동체")
para(
    "마을은 생애주기의 여러 단계에서 훈련과 재학습을 담당하는 공동체다. "
    "산업화 시대에 도시에 기능을 빼앗겨 행정 단위로 전락했던 마을은 "
    "디지털 시대에 '발견된 방향을 훈련으로 잇는 생활권'으로 복원된다. "
    "청소년에게는 마을 멘토단과 훈련 경로가, 중장년에게는 재학습 접근로가 마을을 통해 열린다. "
    "마을은 중앙 정부가 직접 설계하기보다 구성원이 자치적으로 훈련 경로를 만들고 국가가 최소 기준을 보장할 때 가장 잘 작동한다. "
    "센과 누스바움이 강조한 역량—실제로 무언가를 할 수 있는 힘—은 마을 단위의 훈련 환경이 갖춰질 때 비로소 구체적으로 자라난다.",
    indent_first=None)

h2("5.4", "도시 — 발휘의 무대")
para(
    "도시는 산업화 시대에 학교와 공장을 집적하여 표준 인재를 생산·배치하는 기관이었다. "
    "디지털 시대에 도시는 정체성을 실제 기여로 발휘하는 무대이자 생활 기반으로 전환된다. "
    "프로젝트 연결, 직무 전환, 생활 조정이 도시의 기능이 되며, 청년의 첫 기여와 중장년의 재설계가 도시를 통해 이루어진다. "
    "도시에 위임되는 권리의 성격이 '표준 경로에 나를 맞추기'에서 '내 기여를 발휘할 무대 만들기'로 바뀐다.",
    indent_first=None)
para(
    "그러나 모든 기능이 한 대도시에 집중되면 개인이 위임한 권리도 거대한 행정 단위 하나에 묶여 회수하기 어려워진다. "
    "이것은 국가로의 권리 집중과 같은 문제를 도시 규모에서 되풀이하는 것이다. "
    "사람은 자신이 어디에 무엇을 맡겼는지 보이고 필요하면 거두어들일 수 있는 규모 안에서만 권리를 능동적으로 위임한다. "
    "그러므로 디지털 시대의 도시는 거대한 단일 대도시가 아니라, 사람이 체감할 수 있는 규모의 강소도시들이 연합하는 형태로 재편되어야 한다. "
    "강소도시는 한 사람의 강점을 알아보고 기여의 기회를 연결할 만큼 작고 구체적인 생활권이다. "
    "이 규모에서는 내가 맡긴 결정권이 어떻게 쓰이는지 들여다보고 조정할 수 있다. "
    "동시에 작은 도시 하나만으로는 시장과 데이터, 기회의 폭이 부족하므로, 강소도시들은 데이터·인공지능·행정·산업 기반을 "
    "서로 연결해 하나의 연합된 생활권을 이룬다. "
    "발휘의 폭은 연합 전체에서 얻되 권리의 위임과 회수는 내가 속한 작은 도시 단위에서 이루어지는 구조다. "
    "규모의 효율은 연합으로 확보하고 권리의 통제권은 개인 가까이 남기는 것, "
    "이것이 대도시를 강소도시의 연합으로 다시 짜야 하는 이유다.",
    indent_first=None)

h2("5.5", "국가와 세계 — 조율과 공동 위험 관리")
para(
    "국가는 산업화 시대에 권리 집중의 핵심 기구였다. 교육·고용·복지·의료를 중앙에서 일괄 관리했고, "
    "개인이 어느 공동체에 권리를 위임할지 선택할 여지를 거의 남기지 않았다. "
    "디지털 시대의 국가는 모든 것을 관리하는 역할에서 물러나, "
    "가정·이웃·마을·도시가 각자의 역할을 잘 할 수 있도록 최소 조건을 보장하고 이들을 조율하는 "
    "'메타 거버넌스(meta-governance)'로 전환되어야 한다. "
    "모든 사람에게 권리 위임의 기회가 열리도록 최소 권리선을 세우고, "
    "어느 공동체도 권리를 과도하게 독점하지 않도록 규칙을 만드는 것이 국가의 일이다. "
    "에스핑안데르센이 말한 사회적 투자 국가—소득 보전에서 역량 형성으로—는 이 전환의 구체적인 형태다.",
    indent_first=None)
para(
    "세계는 산업화 시대에 국가 간 협정으로만 작동하던 추상적 층위였다. "
    "디지털 시대에 세계는 처음으로 개인의 권리 위임 상대로 등장한다. "
    "국경을 넘는 협업이 가능해지면서 청년과 중장년이 세계 무대에서 기여를 발휘할 수 있다. "
    "동시에 기후변화·인공지능 위험·감염병처럼 국가 단위를 넘어서는 공동 위험이 커지면서 "
    "세계는 모든 생애단계의 안전과 직결된 능동적 거버넌스 층위가 된다. "
    "오스트롬이 만년에 연구한 기후변화 대응 거버넌스가 보여주듯, 거버넌스 층위가 더 많고 더 분산될수록 "
    "공동 위험에 대한 적응력도 높아진다.")

# ════════════════════════ 6. 함의와 결론 ════════════════════════
h1("6.", "디지털 전환이 공동체에 요구하는 것 — 권리의 재분산")

h2("6.1", "권리 분산의 세 원리")
rich([("첫째, 위임의 선택권을 개인에게 돌려주어야 한다. ", True),
      ("국가 하나가 모든 생애단계의 필요를 독점 관리하는 구조에서, 개인이 필요에 따라 가정·이웃·마을·도시에 "
       "권리를 나누어 위임하고 상황이 바뀌면 거두어들일 수 있는 구조로 전환해야 한다. "
       "로크가 말했듯 위임은 항상 회수 가능해야 한다. "
       "탈상품화—시장에 기대지 않고도 살아갈 수 있게 해주는 사회 보장—만으로는 부족하며, "
       "위임의 대상을 스스로 선택하고 바꿀 수 있는 구조가 함께 필요하다.", False)], indent_first=None)
rich([("둘째, 가까운 공동체의 역량을 복원해야 한다. ", True),
      ("브론펜브레너가 지적했듯 개인에게 가장 깊은 영향을 미치는 것은 가까운 환경에서의 직접적 상호작용이다. "
       "산업화 시대에 약화된 이웃과 마을의 역할을 되살리는 것이 권리 재분산의 핵심이다. "
       "이는 중앙 예산을 지방에 내려보내는 것이 아니라, 이웃과 마을이 스스로 역할을 정의하고 "
       "자치적으로 운영할 수 있는 조건을 만드는 것이다. 오스트롬이 보여주었듯 자치는 주어지는 것이 아니라 "
       "구성원들이 스스로 규칙을 만들 때 비로소 작동한다.", False)])
rich([("셋째, 생애주기 단계별 위임 구조를 제도로 명시해야 한다. ", True),
      ("각 생애단계에서 어떤 공동체가 어떤 역할을 맡는지, 개인이 어디에 권리를 위임하는지, "
       "그 위임이 어떤 조건에서 회수될 수 있는지를 사회적 합의로 명시해야 한다. "
       "헥만이 실증한 생애 초기 투자의 복리 효과를 감안하면, 영아기에서 청소년기까지의 공동체 지원은 "
       "복지 지출이 아니라 사회 전체의 역량에 대한 선제적 투자로 재정의되어야 한다.", False)])

h2("6.2", "생애는 누적의 사이클이다")
para(
    "디지털 시대의 생애주기는 직선이 아니라, 앞 단계의 결과가 다음 단계의 조건이 되는 나선형 누적 구조다. "
    "중장년의 재설계는 청소년기의 훈련이 성공했을 때 가능하고, 노년의 전수는 청년·중장년의 발휘가 쌓였을 때 의미를 갖는다. "
    "초기 단계의 권리 위임 격차—어떤 아이는 가정·이웃·마을의 풍부한 지지 속에서 단서를 발견하고, "
    "어떤 아이는 그 공동체들이 모두 약한 환경에서 태어나는—는 생애 전체 사이클의 단절을 의미한다. "
    "헥만의 연구는 이 불평등이 경제학적으로도 가장 비효율적인 낭비임을 보여준다. "
    "일찍 메꾸지 않은 결핍은 나중에 훨씬 큰 사회적 비용을 치르게 만든다.",
    indent_first=None)
para(
    "센과 누스바움이 강조했듯 형식적 권리보다 중요한 것은 실질적 역량이다. "
    "모든 아이에게 학교가 있다고 해서 모든 아이가 발견과 훈련의 역량을 갖는 것은 아니다. "
    "공동체는 그 역량이 실제로 작동하도록 조건을 만들어야 한다. "
    "위임의 구조가 역량의 실현을 뒷받침할 때, 비로소 계약은 정당성을 갖는다.")

h2("6.3", "결론")
para(
    "이 연구의 핵심 주장은 세 줄로 요약된다. 첫째, 공동체는 개인이 자신의 필요에 의해 권리를 위임하는 "
    "계약을 통해 구현되는 실체이며, 위임은 언제든 회수 가능해야 한다. "
    "둘째, 산업화 시대에 권리가 국가로 모인 것은 그 시대 개인의 필요에 따른 선택이었으나, "
    "정체성의 발휘를 원하는 디지털 시대 개인의 필요에는 더 이상 맞지 않는다. "
    "셋째, 디지털 시대의 전환은 기술의 문제가 아니라, 모인 권리를 가정·이웃·마을·도시·국가·세계의 "
    "여러 층위로 다시 분산하고 개인에게 위임의 선택권을 돌려주는 재편의 과제다.",
    indent_first=None)
para(
    "루소가 꿈꾼 자치, 로크가 요구한 신탁의 회수 가능성, 오스트롬이 실증한 다중심 자치 거버넌스, "
    "브론펜브레너가 밝힌 근위 과정의 중요성, 폴라니가 진단한 재배태의 필요—이 모든 논의는 같은 방향을 가리킨다. "
    "개인은 자신의 필요에 맞는 공동체를 스스로 구성하고, 그 공동체에 권리를 위임하며, "
    "상황이 바뀌면 권리를 거두어들일 수 있어야 한다.")

# ════════════════════════ 참고문헌 (APA 7th) ════════════════════════
h1("", "참고문헌")
# 각 항목: (저자·연도 등 정자, 이탤릭 제목/저널, 나머지 정자)
refs = [
    ("Bronfenbrenner, U. (1979). ", "The ecology of human development: Experiments by nature and design", ". Harvard University Press."),
    ("Erikson, E. H. (1950). ", "Childhood and society", ". W. W. Norton."),
    ("Erikson, E. H. (1968). ", "Identity: Youth and crisis", ". W. W. Norton."),
    ("Esping-Andersen, G. (1999). ", "Social foundations of postindustrial economies", ". Oxford University Press."),
    ("Heckman, J. J. (2006). Skill formation and the economics of investing in disadvantaged children. ", "Science, 312", "(5782), 1900–1902. https://doi.org/10.1126/science.1128898"),
    ("Hobbes, T. (1996). ", "Leviathan", " (R. Tuck, Ed.). Cambridge University Press. (Original work published 1651)"),
    ("Locke, J. (1988). ", "Two treatises of government", " (P. Laslett, Ed.). Cambridge University Press. (Original work published 1689)"),
    ("Nussbaum, M. C. (2011). ", "Creating capabilities: The human development approach", ". Harvard University Press."),
    ("Ostrom, E. (1990). ", "Governing the commons: The evolution of institutions for collective action", ". Cambridge University Press."),
    ("Ostrom, E. (2009). ", "A polycentric approach for coping with climate change", " (Policy Research Working Paper No. 5095). World Bank."),
    ("Polanyi, K. (1944). ", "The great transformation: The political and economic origins of our time", ". Farrar & Rinehart."),
    ("Rousseau, J.-J. (1997). ", "The social contract and other later political writings", " (V. Gourevitch, Ed. & Trans.). Cambridge University Press. (Original work published 1762)"),
    ("Sen, A. (1999). ", "Development as freedom", ". Oxford University Press."),
    ("Young, M. (1958). ", "The rise of the meritocracy, 1870–2033", ". Thames and Hudson."),
]
for pre, ital, post in refs:
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(3); pf.line_spacing = 1.4
    pf.left_indent = Mm(8); pf.first_line_indent = Mm(-8)  # hanging indent
    r1 = p.add_run(pre); set_run_font(r1, name=BODY_FONT, size=9.5)
    if ital:
        r2 = p.add_run(ital); set_run_font(r2, name=BODY_FONT, size=9.5, italic=True)
    if post:
        r3 = p.add_run(post); set_run_font(r3, name=BODY_FONT, size=9.5)

add_page_numbers(doc)
os.makedirs(os.path.dirname(OUT), exist_ok=True)
doc.save(OUT)
print("saved:", OUT)
print("paragraphs:", len(doc.paragraphs), "tables:", len(doc.tables))
